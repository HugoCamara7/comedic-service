"""
word_generator.py
------------------
Genera los carnets en Word:

1. Cada persona se renderiza individualmente con docxtpl sobre la plantilla
   de carnet (reemplaza {{ dni }}, {{ nombres_completos }}, {{ foto }}, etc.)
2. El contenido de cada carnet renderizado se copia (con python-docx, a
   nivel de XML) dentro de las celdas de una hoja maestra generada por
   modules/print_layout.py, respetando la grilla filas x columnas
   configurada.
3. Varias hojas se combinan en un unico .docx final con docxcompose.

Esta combinacion evita el problema de "scope" de docxtpl al insertar
subdocumentos con datos distintos por celda: cada carnet se renderiza en su
propio contexto antes de insertarse como contenido ya resuelto.
"""
from __future__ import annotations

import tempfile
from copy import deepcopy
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.shared import Mm
from docxcompose.composer import Composer
from docxtpl import DocxTemplate, InlineImage

from modules.config import PrintConfig
from modules.print_layout import build_master_sheet, carnets_per_page

A_BLIP = qn("a:blip")
R_EMBED = qn("r:embed")


def _build_context(person: Dict, foto_field_present: bool, tpl: DocxTemplate, photo_path: Optional[Path]) -> Dict:
    context = {
        "numero_carnet": person.get("numero_carnet", ""),
        "nombres_completos": person.get("nombre_completo", ""),
        "dni": person.get("dni", ""),
        "fecha_nacimiento": person.get("fecha_nacimiento", ""),
        "fecha_emision": person.get("fecha_emision", ""),
        "fecha_vencimiento": person.get("fecha_vencimiento", ""),
        "empresa": person.get("empresa", ""),
        "puesto": person.get("puesto", ""),
    }
    if foto_field_present:
        if photo_path and photo_path.exists():
            context["foto"] = InlineImage(tpl, str(photo_path), width=Mm(20))
        else:
            context["foto"] = ""
    return context


def render_single_carnet(
    template_path: Path,
    person: Dict,
    output_path: Path,
    foto_field_present: bool,
    photo_path: Optional[Path] = None,
) -> Path:
    """Renderiza un carnet individual con docxtpl y lo guarda en output_path."""
    tpl = DocxTemplate(str(template_path))
    context = _build_context(person, foto_field_present, tpl, photo_path)
    tpl.render(context)
    tpl.save(str(output_path))
    return output_path


def _remap_and_copy_images(element, source_part, dest_part, rid_cache: Dict[str, str]) -> None:
    for blip in element.iter(A_BLIP):
        old_rid = blip.get(R_EMBED)
        if not old_rid:
            continue
        if old_rid not in rid_cache:
            try:
                image_part = source_part.related_parts[old_rid]
            except KeyError:
                continue
            new_rid = dest_part.relate_to(image_part, RT.IMAGE)
            rid_cache[old_rid] = new_rid
        blip.set(R_EMBED, rid_cache[old_rid])


def _copy_body_into_cell(source_doc: Document, dest_cell, dest_doc: Document) -> None:
    tc = dest_cell._tc
    # Limpia el contenido por defecto de la celda (un parrafo vacio).
    for child in list(tc):
        if child.tag in (qn("w:p"), qn("w:tbl")):
            tc.remove(child)

    rid_cache: Dict[str, str] = {}
    body = source_doc.element.body
    for element in list(body):
        if element.tag == qn("w:sectPr"):
            continue
        new_element = deepcopy(element)
        _remap_and_copy_images(new_element, source_doc.part, dest_doc.part, rid_cache)
        tc.append(new_element)


def compose_page(
    people_batch: List[Dict],
    template_path: Path,
    photos_dir: Optional[Path],
    foto_field_present: bool,
    config: PrintConfig,
    work_dir: Path,
) -> Path:
    """Renderiza una hoja (pagina) con hasta filas*columnas carnets."""
    n_slots = carnets_per_page(config)
    master_path = work_dir / "master_sheet.docx"
    build_master_sheet(config, master_path, n_carnets=len(people_batch))
    master_doc = Document(str(master_path))
    table = master_doc.tables[0]

    idx = 0
    for r in range(config.filas):
        for c in range(config.columnas):
            if idx >= len(people_batch):
                break
            person = people_batch[idx]
            photo_path = None
            if foto_field_present and photos_dir and person.get("fotografia"):
                candidate = photos_dir / person["fotografia"]
                if candidate.exists():
                    photo_path = candidate

            individual_path = work_dir / f"_carnet_{idx}.docx"
            render_single_carnet(
                template_path=template_path,
                person=person,
                output_path=individual_path,
                foto_field_present=foto_field_present,
                photo_path=photo_path,
            )
            individual_doc = Document(str(individual_path))
            cell = table.cell(r, c)
            _copy_body_into_cell(individual_doc, cell, master_doc)
            idx += 1

    page_path = work_dir / f"page_{id(people_batch)}.docx"
    master_doc.save(str(page_path))
    return page_path


def generate_word_document(
    people: List[Dict],
    template_path: Path,
    photos_dir: Optional[Path],
    foto_field_present: bool,
    config: PrintConfig,
    output_path: Path,
) -> Path:
    """Genera el documento Word final (todas las hojas combinadas)."""
    n_per_page = carnets_per_page(config)
    batches = [people[i : i + n_per_page] for i in range(0, len(people), n_per_page)]
    if not batches:
        batches = [[]]

    with tempfile.TemporaryDirectory() as tmp:
        work_dir = Path(tmp)
        page_paths = []
        for batch_idx, batch in enumerate(batches):
            batch_dir = work_dir / f"batch_{batch_idx}"
            batch_dir.mkdir(parents=True, exist_ok=True)
            page_path = compose_page(
                people_batch=batch,
                template_path=template_path,
                photos_dir=photos_dir,
                foto_field_present=foto_field_present,
                config=config,
                work_dir=batch_dir,
            )
            page_paths.append(page_path)

        base_doc = Document(str(page_paths[0]))
        composer = Composer(base_doc)
        for extra_page in page_paths[1:]:
            base_doc.add_page_break()
            composer.append(Document(str(extra_page)))

        output_path.parent.mkdir(parents=True, exist_ok=True)
        composer.save(str(output_path))

    return output_path
