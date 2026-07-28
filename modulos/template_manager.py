"""
template_manager.py
--------------------
Gestion de plantillas Word (carnet individual).

- Genera una plantilla por defecto (util para probar la aplicacion sin subir
  nada) usando python-docx, con los campos jinja que espera docxtpl.
- Permite guardar/listar plantillas subidas por el usuario en templates/.
- Detecta si una plantilla usa el campo {{ foto }} (para exigir fotografia).

Campos jinja soportados (deben existir literalmente en la plantilla, en
parrafos o dentro de tablas):
    {{ numero_carnet }}
    {{ nombres_completos }}
    {{ dni }}
    {{ fecha_nacimiento }}
    {{ fecha_emision }}
    {{ fecha_vencimiento }}
    {{ empresa }}
    {{ puesto }}
    {{ foto }}
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import List, Set

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = BASE_DIR / "templates"
TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

DEFAULT_TEMPLATE_NAME = "plantilla_default.docx"

KNOWN_FIELDS = [
    "numero_carnet",
    "nombres_completos",
    "dni",
    "fecha_nacimiento",
    "fecha_emision",
    "fecha_vencimiento",
    "empresa",
    "puesto",
    "foto",
]

MAX_TEMPLATE_SIZE_MB = 15
ALLOWED_TEMPLATE_EXTENSIONS = {".docx"}


class TemplateError(Exception):
    pass


def _set_cell_background(cell, color_hex: str) -> None:
    shd = cell._tc.get_or_add_tcPr()
    shading = shd.makeelement(qn("w:shd"), {qn("w:fill"): color_hex})
    shd.append(shading)


def _no_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{edge}"), {qn("w:val"): "none"})
        borders.append(el)
    tbl_pr.append(borders)


def build_default_template(
    path: Path,
    ancho_mm: float = 85.6,
    alto_mm: float = 54.0,
) -> Path:
    """Crea una plantilla de carnet por defecto (tamano tipo tarjeta) con los
    campos jinja esperados por docxtpl, distribuidos en una tabla de 2
    columnas (foto | datos)."""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Mm(ancho_mm)
    section.page_height = Mm(alto_mm)
    section.top_margin = Mm(2)
    section.bottom_margin = Mm(2)
    section.left_margin = Mm(2)
    section.right_margin = Mm(2)

    usable_width = ancho_mm - 4
    foto_col_width = min(20.0, usable_width * 0.32)
    datos_col_width = usable_width - foto_col_width

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_borders(table)
    table.columns[0].width = Mm(foto_col_width)
    table.columns[1].width = Mm(datos_col_width)
    row = table.rows[0]
    row.cells[0].width = Mm(foto_col_width)
    row.cells[1].width = Mm(datos_col_width)

    # --- Celda de foto ---
    foto_cell = row.cells[0]
    foto_cell.text = ""
    p = foto_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{ foto }}")
    run.font.size = Pt(7)

    # --- Celda de datos ---
    datos_cell = row.cells[1]
    datos_cell.text = ""
    first = True

    def add_line(text: str, bold_label: str = "", size: int = 8, color=None):
        nonlocal first
        par = datos_cell.paragraphs[0] if first else datos_cell.add_paragraph()
        first = False
        par.paragraph_format.space_after = Pt(1)
        if bold_label:
            r1 = par.add_run(bold_label)
            r1.bold = True
            r1.font.size = Pt(size)
        r2 = par.add_run(text)
        r2.font.size = Pt(size)
        if color:
            r2.font.color.rgb = RGBColor(*color)

    add_line("CARNET DE SANIDAD", size=9)
    doc_title_par = datos_cell.paragraphs[0]
    doc_title_par.runs[0].bold = True
    doc_title_par.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    add_line("{{ nombres_completos }}", "Nombre: ", size=8)
    add_line("{{ dni }}", "DNI: ", size=8)
    add_line("{{ puesto }}", "Puesto: ", size=8)
    add_line("{{ empresa }}", "Empresa: ", size=8)
    add_line("{{ numero_carnet }}", "N Carnet: ", size=8)
    add_line("{{ fecha_emision }}", "Emision: ", size=7)
    add_line("{{ fecha_vencimiento }}", "Vence: ", size=7)

    doc.save(path)
    return path


def ensure_default_template() -> Path:
    path = TEMPLATES_DIR / DEFAULT_TEMPLATE_NAME
    if not path.exists():
        build_default_template(path)
    return path


def list_templates() -> List[Path]:
    return sorted(TEMPLATES_DIR.glob("*.docx"))


def validate_extension(filename: str) -> None:
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_TEMPLATE_EXTENSIONS:
        raise TemplateError(f"Formato de plantilla no permitido: '{ext}'. Solo se acepta .docx")


def save_uploaded_template(file_like, filename: str, max_size_mb: int = MAX_TEMPLATE_SIZE_MB) -> Path:
    validate_extension(filename)
    data = file_like.read() if hasattr(file_like, "read") else Path(file_like).read_bytes()
    size_mb = len(data) / (1024 * 1024)
    if size_mb > max_size_mb:
        raise TemplateError(f"La plantilla pesa {size_mb:.1f} MB y excede el limite de {max_size_mb} MB.")

    safe_name = re.sub(r"[^A-Za-z0-9._-]", "_", Path(filename).name)
    dest = TEMPLATES_DIR / safe_name
    dest.write_bytes(data)

    # Validar que el docx se pueda abrir y contenga al menos un campo conocido.
    try:
        fields = detect_fields(dest)
    except Exception as exc:  # noqa: BLE001
        dest.unlink(missing_ok=True)
        raise TemplateError(f"El archivo no es un .docx valido: {exc}") from exc

    if not fields:
        dest.unlink(missing_ok=True)
        raise TemplateError(
            "La plantilla no contiene ningun campo reconocido "
            f"({', '.join('{{ ' + f + ' }}' for f in KNOWN_FIELDS)})."
        )

    return dest


def _iter_all_text(doc: Document):
    for p in doc.paragraphs:
        yield p.text
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p.text


def detect_fields(path: Path) -> Set[str]:
    """Devuelve el conjunto de campos jinja conocidos presentes en la plantilla."""
    doc = Document(str(path))
    full_text = "\n".join(_iter_all_text(doc))
    found = set()
    for field in KNOWN_FIELDS:
        if re.search(r"\{\{\s*" + re.escape(field) + r"\s*\}\}", full_text):
            found.add(field)
    return found


def template_requires_photo(path: Path) -> bool:
    return "foto" in detect_fields(path)
