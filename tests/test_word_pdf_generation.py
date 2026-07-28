from pathlib import Path

import pytest
from docx import Document

from modules import pdf_generator, template_manager, word_generator
from modules.config import AppConfig


@pytest.fixture
def template_path(tmp_path):
    path = tmp_path / "plantilla.docx"
    template_manager.build_default_template(path)
    return path


def _person(i: int) -> dict:
    return {
        "dni": f"1000000{i}",
        "nombre_completo": f"Persona {i} Prueba",
        "numero_carnet": f"C-{i:05d}",
        "fecha_nacimiento": "01/01/1990",
        "fecha_emision": "01/01/2026",
        "fecha_vencimiento": "01/01/2027",
        "empresa": "Empresa Test",
        "puesto": "Puesto Test",
        "fotografia": "",
    }


def test_detect_fields_plantilla_default(template_path):
    fields = template_manager.detect_fields(template_path)
    assert "dni" in fields
    assert "nombres_completos" in fields
    assert "foto" in fields


def test_render_single_carnet(tmp_path, template_path):
    out = tmp_path / "carnet.docx"
    word_generator.render_single_carnet(
        template_path=template_path,
        person=_person(1),
        output_path=out,
        foto_field_present=False,
    )
    assert out.exists()
    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    assert "Persona 1 Prueba" in full_text
    assert "10000001" in full_text


def test_generacion_word_una_persona(tmp_path, template_path):
    cfg = AppConfig()
    cfg.print.columnas = 2
    cfg.print.filas = 2
    out = tmp_path / "salida.docx"
    word_generator.generate_word_document(
        people=[_person(1)],
        template_path=template_path,
        photos_dir=None,
        foto_field_present=False,
        config=cfg.print,
        output_path=out,
    )
    assert out.exists()
    doc = Document(str(out))
    assert len(doc.tables) == 1  # una sola hoja


def test_generacion_word_varias_personas(tmp_path, template_path):
    cfg = AppConfig()
    cfg.print.columnas = 2
    cfg.print.filas = 2  # 4 por hoja
    out = tmp_path / "salida.docx"
    people = [_person(i) for i in range(1, 5)]
    word_generator.generate_word_document(
        people=people,
        template_path=template_path,
        photos_dir=None,
        foto_field_present=False,
        config=cfg.print,
        output_path=out,
    )
    assert out.exists()
    doc = Document(str(out))
    assert len(doc.tables) == 1  # exactamente llena una hoja


def test_generacion_word_ultima_pagina_incompleta(tmp_path, template_path):
    cfg = AppConfig()
    cfg.print.columnas = 2
    cfg.print.filas = 2  # 4 por hoja
    out = tmp_path / "salida.docx"
    people = [_person(i) for i in range(1, 6)]  # 5 personas -> 2 hojas, la 2a con 1
    word_generator.generate_word_document(
        people=people,
        template_path=template_path,
        photos_dir=None,
        foto_field_present=False,
        config=cfg.print,
        output_path=out,
    )
    doc = Document(str(out))
    assert len(doc.tables) == 2  # dos hojas maestras

    def _all_text(document):
        texts = []

        def walk_cell(cell):
            for p in cell.paragraphs:
                texts.append(p.text)
            for t in cell.tables:
                for row in t.rows:
                    for c in row.cells:
                        walk_cell(c)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    walk_cell(cell)
        return "\n".join(texts)

    full_text = _all_text(doc)
    for i in range(1, 6):
        assert f"Persona {i} Prueba" in full_text


def test_generacion_pdf(tmp_path, template_path):
    cfg = AppConfig()
    cfg.print.columnas = 2
    cfg.print.filas = 2
    out = tmp_path / "salida.docx"
    word_generator.generate_word_document(
        people=[_person(1), _person(2)],
        template_path=template_path,
        photos_dir=None,
        foto_field_present=False,
        config=cfg.print,
        output_path=out,
    )
    pdf_path = pdf_generator.convert_docx_to_pdf(out, tmp_path)
    assert pdf_path.exists()
    assert pdf_path.suffix == ".pdf"
    assert pdf_path.stat().st_size > 0


def test_foto_faltante_no_rompe_generacion(tmp_path, template_path):
    """Una foto inexistente en disco no debe hacer fallar la generacion;
    el campo simplemente queda vacio (la validacion previa es la que decide
    si el registro es correcto u observado)."""
    cfg = AppConfig()
    cfg.print.columnas = 2
    cfg.print.filas = 1
    out = tmp_path / "salida.docx"
    person = _person(1)
    person["fotografia"] = "no_existe.jpg"
    word_generator.generate_word_document(
        people=[person],
        template_path=template_path,
        photos_dir=tmp_path,  # no existe el archivo dentro
        foto_field_present=True,
        config=cfg.print,
        output_path=out,
    )
    assert out.exists()
